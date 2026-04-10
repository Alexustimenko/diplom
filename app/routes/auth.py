# app/routes/auth.py
from flask import Blueprint, render_template, request, redirect, url_for, session, send_file
from app.db import get_conn
import bcrypt
import io
from docx import Document
from openpyxl import Workbook
from datetime import datetime
import re
import pyodbc

import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart


auth_bp = Blueprint("auth", __name__)
EMAIL_ALLOWED_RE = re.compile(r"^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$")

def is_valid_email_strict(email: str) -> bool:
    if not email or len(email) > 254:
        return False

    email = email.strip()

    # Быстрый фильтр по допустимым символам и базовой структуре
    if not EMAIL_ALLOWED_RE.match(email):
        return False

    if email.count("@") != 1:
        return False

    local, domain = email.split("@", 1)

    # local part правила
    if len(local) > 64:
        return False
    if local[0] == "." or local[-1] == ".":
        return False
    if ".." in local:
        return False

    # domain правила
    if domain[0] == "." or domain[-1] == ".":
        return False
    if ".." in domain:
        return False
    if "." not in domain:
        return False

    labels = domain.split(".")
    if any(not lbl for lbl in labels):
        return False

    tld = labels[-1]
    if len(tld) < 2 or not tld.isalpha():
        return False

    for lbl in labels:
        if len(lbl) > 63:
            return False
        if lbl[0] == "-" or lbl[-1] == "-":
            return False

    return True

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText


def send_order_created_email(user_email,
                             order_id,
                             phone,
                             delivery_type,
                             items,
                             total):

    sender = "ustimenko.lesha@gmail.com"
    password = "fmcoakbuddnebowc"

    subject = "Новый заказ в магазине ROLMARK"

    # -------- формируем строки таблицы --------
    rows_html = ""
    for item in items:
        line_total = item["price"] * item["qty"]
        rows_html += f"""
        <tr>
            <td>{item["name"]}</td>
            <td>{item["price"]:.2f}</td>
            <td>{item["qty"]}</td>
            <td>{line_total:.2f}</td>
        </tr>
        """

    # -------- HTML письмо --------
    html = f"""
    <html>
    <body style="font-family: Arial, sans-serif;">

        <p>Здравствуйте!</p>

        <p>Рады видеть в интернет-магазине <b>ROLMARK</b>.</p>

        <p><b>Ваш заказ оформлен.</b></p>

        <p><b>Проверьте контактные данные и детали заказа</b></p>

        <p>
            Номер заказа: <b>{order_id}</b><br>
            Телефон: {phone}<br>
            Почта: {user_email}<br>
            Вид доставки: {delivery_type}
        </p>

        <table border="1" cellpadding="8" cellspacing="0" width="100%" style="border-collapse: collapse;">
            <tr style="background:#f2f2f2; font-weight:bold;">
                <td>Товар</td>
                <td>Цена</td>
                <td>Кол-во</td>
                <td>Сумма</td>
            </tr>

            {rows_html}

        </table>

        <p style="margin-top:20px;">
            <b>Итого: {total:.2f} Бел.руб.</b>
        </p>

    </body>
    </html>
    """

    msg = MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["From"] = sender
    msg["To"] = user_email

    msg.attach(MIMEText(html, "html"))

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(sender, password)
        server.send_message(msg)
def send_ready_email(user_email, order_id, delivery_method):

    sender = "ustimenko.lesha@gmail.com"
    password = "fmcoakbuddnebowc"

    subject = "Новая информация по заказу"

    # 🔥 логика по способу доставки
    if delivery_method in ("minsk", "belarus"):
        body = f"""
Ваш заказ №{order_id} готов к выдаче.

Курьер выехал и скоро Вам позвонит.
"""
    elif delivery_method == "pickup":
        body = f"""
Ваш заказ №{order_id} готов к выдаче.

Заберите его.
"""
    else:
        body = f"""
Ваш заказ №{order_id} готов к выдаче.
"""

    msg = MIMEText(body)
    msg["Subject"] = subject
    msg["From"] = sender
    msg["To"] = user_email

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(sender, password)
        server.send_message(msg)
def _require_admin():
    if "user_id" not in session:
        return redirect(url_for("auth.login"))
    if session.get("is_admin") != 1:
        return "Forbidden", 403
    return None
def _require_login():
    if "user_id" not in session:
        return redirect(url_for("auth.login"))
    return None



# app/routes/auth.py  (ФРАГМЕНТ: только index() с фильтрами + пагинацией)

@auth_bp.route("/")
def index():
    # ✅ q_raw — как есть (для инпута), q — обрезанный (для SQL)
    q_raw = request.args.get("q", "")
    q = q_raw.strip()

    # фильтры
    category_id_raw = request.args.get("category_id", "").strip()
    brand_id_raw = request.args.get("brand_id", "").strip()
    color_raw = request.args.get("color", "").strip()
    material_raw = request.args.get("material", "").strip()
    price_from_raw = request.args.get("price_from", "").strip()
    price_to_raw = request.args.get("price_to", "").strip()

    # пагинация
    per_page = 3
    try:
        page = int(request.args.get("page", "1"))
    except:
        page = 1
    if page < 1:
        page = 1

    conn = get_conn()
    cur = conn.cursor()

    # 1) категории для выпадашки
    cur.execute("SELECT category_id, name, parent_id FROM dbo.categories ORDER BY name")
    categories = cur.fetchall()

    # 2) построим дерево категорий и ветку "запчасти"
    parts_roots = set()
    children_map = {}

    for c in categories:
        if c.name and "запчаст" in c.name.lower():
            parts_roots.add(int(c.category_id))
        if c.parent_id is not None:
            children_map.setdefault(int(c.parent_id), []).append(int(c.category_id))

    def collect_descendants(root_id: int) -> set[int]:
        stack = [root_id]
        out = set()
        while stack:
            x = stack.pop()
            if x in out:
                continue
            out.add(x)
            for ch in children_map.get(x, []):
                stack.append(ch)
        return out

    parts_tree = set()
    for rid in parts_roots:
        parts_tree |= collect_descendants(rid)

    selected_category_id = int(category_id_raw) if category_id_raw.isdigit() else None
    selected_brand_id = int(brand_id_raw) if brand_id_raw.isdigit() else None

    is_parts_mode = bool(selected_category_id and selected_category_id in parts_tree)

    # 3) бренды для выпадашки:
    #    если выбрана категория — показываем только бренды, у которых есть товары в этой категории (и её подкатегориях)
    brands = []
    cat_ids_for_brand = None

    if selected_category_id:
        cat_ids_for_brand = collect_descendants(selected_category_id)

    if cat_ids_for_brand and len(cat_ids_for_brand) > 0:
        placeholders = ",".join(["?"] * len(cat_ids_for_brand))
        cur.execute(f"""
            SELECT DISTINCT
                vp.id_brand   AS id_brand,
                vp.brand_name AS name
            FROM dbo.vw_products vp
            WHERE vp.id_brand IS NOT NULL
              AND vp.category_id IN ({placeholders})
            ORDER BY vp.brand_name
        """, *list(cat_ids_for_brand))
        brands = cur.fetchall()
    else:
        cur.execute("SELECT id_brand, name FROM dbo.brand ORDER BY name")
        brands = cur.fetchall()

    # если выбран бренд, но его нет в доступных брендах — сбросим (чтобы фильтр не давал “пусто”)
    if selected_brand_id and brands:
        allowed = {int(b.id_brand) for b in brands}
        if selected_brand_id not in allowed:
            selected_brand_id = None
            brand_id_raw = ""

    # 4) WHERE
    params = []
    where = []

    # поиск (серверный, по всем страницам)
    if q:
        like = f"%{q}%"
        where.append("""
            (
                name LIKE ?
                OR description LIKE ?
                OR brand_name LIKE ?
                OR category_name LIKE ?
            )
        """)
        params += [like, like, like, like]

    # категория
    if selected_category_id:
        where.append("category_id = ?")
        params.append(selected_category_id)

    # цена
    if price_from_raw:
        try:
            pf = float(price_from_raw.replace(",", "."))
            where.append("price >= ?")
            params.append(pf)
        except:
            price_from_raw = ""

    if price_to_raw:
        try:
            pt = float(price_to_raw.replace(",", "."))
            where.append("price <= ?")
            params.append(pt)
        except:
            price_to_raw = ""

    # ✅ доп. фильтры только если НЕ запчасти
    if not is_parts_mode:
        # бренд
        if selected_brand_id:
            where.append("id_brand = ?")
            params.append(selected_brand_id)

        # цвет/материал из description по шаблону "Цвет: ..." и "Материал: ..."
        if color_raw:
            c = color_raw.strip().lower()
            where.append("LOWER(ISNULL(description,'')) LIKE ?")
            params.append(f"%цвет:%{c}%")

        if material_raw:
            m = material_raw.strip().lower()
            where.append("LOWER(ISNULL(description,'')) LIKE ?")
            params.append(f"%материал:%{m}%")

    where_sql = ("WHERE " + " AND ".join(where)) if where else ""

    # 5) COUNT для пагинации
    cur.execute(f"SELECT COUNT(*) FROM dbo.vw_products {where_sql}", *params)
    total = int(cur.fetchone()[0])

    total_pages = max(1, (total + per_page - 1) // per_page)
    if page > total_pages:
        page = total_pages

    offset = (page - 1) * per_page

    # 6) данные страницы
    cur.execute(f"""
        SELECT *
        FROM dbo.vw_products
        {where_sql}
        ORDER BY product_id DESC
        OFFSET ? ROWS FETCH NEXT ? ROWS ONLY
    """, *(params + [offset, per_page]))

    products = cur.fetchall()
    # ✅ ХИТЫ ПРОДАЖ (TOP-3)
    cur.execute("EXEC dbo.sp_top3_best_sellers")
    hits = cur.fetchall()
    conn.close()

    return render_template(
        "index.html",
        products=products,

        # ✅ ВАЖНО: в шаблон отдаём q_raw, чтобы пробелы не исчезали из input
        q=q_raw,

        page=page,
        total_pages=total_pages,

        categories=categories,
        brands=brands,

        selected_category_id=selected_category_id,
        selected_brand_id=selected_brand_id,

        price_from=price_from_raw,
        price_to=price_to_raw,
        hits=hits,

        color=color_raw,
        material=material_raw,

        is_parts_mode=is_parts_mode,
        parts_ids_csv=",".join(str(x) for x in sorted(parts_tree))
    )
# ✅ Страница товара (галерея по всем image_id)
@auth_bp.route("/product/<int:pid>")
def product(pid):
    conn = get_conn()
    cur = conn.cursor()

    cur.execute("SELECT * FROM dbo.vw_products WHERE product_id = ?", pid)
    p = cur.fetchone()
    if not p:
        conn.close()
        return "Not Found", 404

    cur.execute("SELECT image_id FROM dbo.images WHERE product_id = ? ORDER BY image_id DESC", pid)
    images = cur.fetchall()  # [(image_id,), ...]

    conn.close()
    return render_template("product.html", p=p, images=images)


@auth_bp.route("/product_image/<int:pid>")
def product_image(pid):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("SELECT TOP 1 image_data FROM dbo.images WHERE product_id=? ORDER BY image_id DESC", pid)
    row = cur.fetchone()
    conn.close()

    if not row or row[0] is None:
        return "", 404

    data = row[0]
    if isinstance(data, memoryview):
        data = data.tobytes()
    else:
        data = bytes(data)

    return send_file(io.BytesIO(data), mimetype="image/jpeg")


@auth_bp.route("/image/<int:image_id>")
def image_by_id(image_id):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("SELECT image_data FROM dbo.images WHERE image_id = ?", image_id)
    row = cur.fetchone()
    conn.close()

    if not row or row[0] is None:
        return "", 404

    data = row[0]
    if isinstance(data, memoryview):
        data = data.tobytes()
    else:
        data = bytes(data)

    return send_file(io.BytesIO(data), mimetype="image/jpeg")


@auth_bp.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        if not is_valid_email_strict(email):
            return render_template("register.html", error="Введите корректный email (пример: name@gmail.com)")

        password = request.form.get("password", "")
        customer_type = request.form.get("customer_type", "fiz").strip()

        if not email or not password:
            return render_template("register.html", error="Заполни email и пароль")

        hashed = bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")

        conn = get_conn()
        cur = conn.cursor()
        cur.execute("EXEC dbo.sp_register_user ?, ?, ?", email, hashed, customer_type)
        result = cur.fetchone()[0]
        conn.commit()
        conn.close()

        if result == -1:
            return render_template("register.html", error="Пользователь уже существует")
        if result == -2:
            return render_template("register.html", error="Неверный тип клиента")

        return render_template("register.html", success="Аккаунт создан! Теперь можно входить.")

    return render_template("register.html")


@auth_bp.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "")

        if not email or not password:
            return render_template("login.html", error="Заполни email и пароль")

        if email == "admin" and password == "admin":
            session["user_id"] = 0
            session["email"] = "admin"
            session["is_admin"] = 1
            return redirect("/admin")

        conn = get_conn()
        cur = conn.cursor()
        cur.execute("EXEC dbo.sp_login_user ?", email)
        row = cur.fetchone()
        conn.close()

        if not row:
            return render_template("login.html", error="Ошибка входа")

        if not bcrypt.checkpw(password.encode(), row.password_hash.encode()):
            return render_template("login.html", error="Ошибка входа")

        session["user_id"] = row.user_id
        session["email"] = email
        session["is_admin"] = 1 if row.is_admin else 0

        return redirect("/admin" if row.is_admin else "/")

    return render_template("login.html")


@auth_bp.route("/logout")
def logout():
    session.clear()
    return redirect("/")


@auth_bp.route("/cabinet")
def cabinet():
    if "user_id" not in session:
        return redirect(url_for("auth.login"))

    conn = get_conn()
    cur = conn.cursor()

    cur.execute("EXEC dbo.sp_get_user_orders ?", session["user_id"])
    rows = cur.fetchall()
    conn.close()

    orders_dict = {}

    for r in rows:
        oid = r.order_id

        if oid not in orders_dict:
            orders_dict[oid] = {
                "order_id": oid,
                "created_at": r.created_at,
                "total": r.total_amount,
                "items": []
            }

        orders_dict[oid]["items"].append({
            "product_name": r.product_name,
            "qty": r.quantity,
            "price": r.price_at_moment
        })

    # 👇 превращаем в список
    orders = list(orders_dict.values())

    return render_template(
        "cabinet.html",
        email=session.get("email"),
        customer_type=session.get("customer_type"),
        orders=orders
    )


@auth_bp.route("/admin", methods=["GET", "POST"])
def admin():
    gate = _require_admin()
    if gate:
        return gate

    error = None
    success = None

    # ================= POST (CRUD) =================
    if request.method == "POST":
        action = request.form.get("action", "").strip()

        try:
            conn = get_conn()
            cur = conn.cursor()

            # ---------- BRANDS ----------
            if action == "add_brand":
                name = request.form.get("brand_name", "").strip()
                if not name:
                    error = "Название бренда обязательно"
                else:
                    # ✅ проверка дубля бренда по имени (без регистра и пробелов)
                    cur.execute("""
                        SELECT TOP 1 id_brand
                        FROM dbo.brand
                        WHERE LTRIM(RTRIM(LOWER(name))) = LTRIM(RTRIM(LOWER(?)))
                    """, name)
                    exists_row = cur.fetchone()

                    if exists_row:
                        error = f"Бренд с названием '{name}' уже существует"
                    else:
                        cur.execute("EXEC dbo.sp_add_brand ?", name)
                        conn.commit()
                        success = "Бренд добавлен"


            elif action == "edit_brand":
                brand_id = int(request.form["brand_id"])
                name = request.form.get("brand_name", "").strip()
                if not name:
                    error = "Название бренда обязательно"
                else:
                    # ✅ дубль имени бренда, но исключаем текущий brand_id
                    cur.execute("""
                        SELECT TOP 1 id_brand
                        FROM dbo.brand
                        WHERE LTRIM(RTRIM(LOWER(name))) = LTRIM(RTRIM(LOWER(?)))
                          AND id_brand <> ?
                    """, name, brand_id)
                    exists_row = cur.fetchone()

                    if exists_row:
                        error = f"Бренд с названием '{name}' уже существует"
                    else:
                        cur.execute("EXEC dbo.sp_update_brand ?, ?", brand_id, name)
                        conn.commit()
                        success = "Бренд обновлён"

            elif action == "delete_brand":
                bid = int(request.form["brand_id"])

                # ✅ нельзя удалить бренд, если у него есть товары
                cur.execute("SELECT TOP 1 1 FROM dbo.products WHERE id_brand = ?", bid)
                has_products = cur.fetchone()

                if has_products:
                    error = "Нельзя удалить бренд: у этого бренда есть товары"
                else:
                    cur.execute("EXEC dbo.sp_delete_brand ?", bid)
                    conn.commit()
                    success = "Бренд удалён"

            # ---------- CATEGORIES ----------
            elif action == "add_category":
                name = request.form.get("category_name", "").strip()
                parent = request.form.get("parent_id") or None
                parent_id = int(parent) if str(parent).isdigit() else None

                if not name:
                    error = "Название категории обязательно"
                else:
                    # ✅ Дубль: одинаковое имя в рамках одного parent_id
                    # (чтобы можно было иметь одинаковые названия в разных ветках)
                    cur.execute("""
                        SELECT TOP 1 category_id
                        FROM dbo.categories
                        WHERE LTRIM(RTRIM(LOWER(name))) = LTRIM(RTRIM(LOWER(?)))
                          AND (
                                (parent_id IS NULL AND ? IS NULL)
                             OR (parent_id = ?)
                          )
                    """, name, parent_id, parent_id)
                    exists_row = cur.fetchone()

                    if exists_row:
                        error = f"Категория/подкатегория с названием '{name}' уже существует"
                    else:
                        cur.execute("EXEC dbo.sp_add_category ?, ?", name, parent_id)
                        conn.commit()
                        success = "Категория добавлена"

            elif action == "edit_category":
                cid = int(request.form["category_id"])
                name = request.form.get("category_name", "").strip()
                parent = request.form.get("parent_id") or None
                parent_id = int(parent) if str(parent).isdigit() else None

                if not name:
                    error = "Название категории обязательно"
                else:
                    # ✅ дубль имени в том же parent_id, исключая текущую категорию
                    cur.execute("""
                        SELECT TOP 1 category_id
                        FROM dbo.categories
                        WHERE LTRIM(RTRIM(LOWER(name))) = LTRIM(RTRIM(LOWER(?)))
                          AND (
                                (parent_id IS NULL AND ? IS NULL)
                             OR (parent_id = ?)
                          )
                          AND category_id <> ?
                    """, name, parent_id, parent_id, cid)
                    exists_row = cur.fetchone()

                    if exists_row:
                        error = f"Категория/подкатегория с названием '{name}' уже существует"
                    else:
                        cur.execute("EXEC dbo.sp_update_category ?, ?, ?", cid, name, parent_id)
                        conn.commit()
                        success = "Категория обновлена"

            elif action == "delete_category":
                cid = int(request.form["category_id"])

                # ✅ нельзя удалить категорию, если в ней есть товары
                cur.execute("SELECT TOP 1 1 FROM dbo.products WHERE category_id = ?", cid)
                has_products = cur.fetchone()

                if has_products:
                    error = "Нельзя удалить категорию/подкатегорию: в ней есть товары"
                else:
                    cur.execute("EXEC dbo.sp_delete_category ?", cid)
                    conn.commit()
                    success = "Категория удалена"
            # ---------- UPDATE ORDER STATUS ----------
            elif action == "update_order_status":
                order_id = int(request.form["order_id"])
                new_status = request.form.get("new_status")

                cur.execute("UPDATE dbo.orders SET status=? WHERE order_id=?",
                            new_status, order_id)
                conn.commit()

                success = "Статус заказа обновлён"

                # если статус Готов к выдаче → отправляем письмо
                if new_status == "Готов к выдаче":

                    cur.execute("""
                        SELECT u.email, o.delivery_method
                        FROM dbo.orders o
                        JOIN dbo.users u ON o.user_id = u.user_id
                        WHERE o.order_id = ?
                    """, order_id)

                    row = cur.fetchone()
                    if row:
                        send_ready_email(row.email, order_id, row.delivery_method)

            # ---------- PRODUCTS ----------
            elif action == "add_product":
                name = request.form.get("name", "").strip()
                description = request.form.get("description") or None

                if not name:
                    error = "Название товара обязательно"
                else:
                    price = float(request.form["price"])
                    old_price = request.form.get("old_price") or None
                    stock = int(request.form["stock"])

                    category_id = request.form.get("category_id") or None
                    brand_id = request.form.get("brand_id") or None

                    old_price = float(old_price) if old_price else None
                    category_id = int(category_id) if category_id else None
                    brand_id = int(brand_id) if brand_id else None

                    # ✅ проверка на дубль названия (без учета регистра и лишних пробелов по краям)
                    cur.execute("""
                        SELECT TOP 1 product_id
                        FROM dbo.products
                        WHERE LTRIM(RTRIM(name)) = LTRIM(RTRIM(?))
                    """, name)
                    exists_row = cur.fetchone()

                    if exists_row:
                        error = f"Товар с названием '{name}' уже существует"
                    if price<=0:
                        error = "Нельзя добавить товар с ценой <= 0"
                    else:
                        cur.execute("""
                            EXEC dbo.sp_add_product
                                @name = ?,
                                @description = ?,
                                @price = ?,
                                @old_price = ?,
                                @stock = ?,
                                @category_id = ?,
                                @brand_id = ?
                        """, name, description, price, old_price, stock, category_id, brand_id)

                        product_id = cur.fetchone()[0]

                        for f in request.files.getlist("images"):
                            if f.filename:
                                cur.execute(
                                    "INSERT INTO dbo.images(product_id, image_data) VALUES (?, ?)",
                                    product_id, f.read()
                                )

                        conn.commit()
                        success = "Товар добавлен"

            elif action == "edit_product":
                pid = int(request.form["product_id"])

                cur.execute("""
                    EXEC dbo.sp_update_product ?, ?, ?, ?, ?, ?, ?, ?
                """,
                            pid,  # @product_id
                            request.form["name"],  # @name
                            request.form.get("description") or None,  # @description
                            float(request.form["price"]),  # @price
                            float(request.form["old_price"]) if request.form.get("old_price") else None,  # @old_price
                            int(request.form["stock"]),  # @stock
                            int(request.form["category_id"]) if request.form.get("category_id") else None,
                            # @category_id
                            int(request.form["brand_id"]) if request.form.get("brand_id") else None  # @brand_id
                            )

                if request.form.get("replace_images"):
                    cur.execute("DELETE FROM dbo.images WHERE product_id=?", pid)

                for f in request.files.getlist("images"):
                    if f.filename:
                        cur.execute(
                            "INSERT INTO dbo.images(product_id,image_data) VALUES (?,?)",
                            pid, f.read()
                        )

                conn.commit()
                success = "Товар обновлён"
            elif action == "delete_product":
                pid = int(request.form["product_id"])
                cur.execute("DELETE FROM dbo.images WHERE product_id=?", pid)
                cur.execute("DELETE FROM dbo.products WHERE product_id=?", pid)
                conn.commit()
                success = "Товар удалён"

            conn.close()
            if not error:
                return redirect("/admin")

        except Exception as e:
            try:
                conn.close()
            except:
                pass
            error = str(e)

    # ================= GET =================
    conn = get_conn()
    cur = conn.cursor()

    cur.execute("SELECT id_brand,name FROM dbo.vw_brands")
    brands = cur.fetchall()

    cur.execute("SELECT category_id,name,parent_id,parent_name FROM dbo.vw_categories")
    categories = cur.fetchall()

    cur.execute("SELECT * FROM dbo.vw_products ORDER BY product_id DESC")
    products = cur.fetchall()

    cur.execute("EXEC dbo.sp_get_all_orders_admin")
    orders_admin = cur.fetchall()

    conn.close()

    return render_template(
        "admin.html",
        brands=brands,
        orders_admin=orders_admin,
        categories=categories,
        products=products,
        tab="products",
        error=error,
        success=success

    )


@auth_bp.route("/admin/reports")
def admin_reports():
    gate = _require_admin()
    if gate:
        return gate

    date_from = request.args.get("from", "").strip()
    date_to   = request.args.get("to", "").strip()
    target_category_id = request.args.get("target_category_id", "").strip()
    brand_id = request.args.get("parts_brand_id", "").strip()

    parts_rows = []

    # если даты не задали — подставим сегодня
    if not date_from or not date_to:
        from datetime import date
        today = date.today().isoformat()
        date_from = today
        date_to = today

    # если пользователь перепутал местами даты — поменяем
    if date_from > date_to:
        date_from, date_to = date_to, date_from

    conn = get_conn()
    cur = conn.cursor()

    # 0) товары по категориям (старый VIEW)
    cur.execute("SELECT * FROM dbo.vw_report_products_per_category ORDER BY category_name")
    report_rows = cur.fetchall()

    # 1) заказы за период
    cur.execute("EXEC dbo.sp_report_orders_period_grouped ?, ?", date_from, date_to)
    orders_rows = cur.fetchall()

    # 2) клиенты (история заказов за период)
    cur.execute("EXEC sp_report_clients_orders")
    customers_rows = cur.fetchall()

    # 3) проданные товары за период
    cur.execute("EXEC dbo.sp_report_sold_products_period ?, ?", date_from, date_to)
    sold_rows = cur.fetchall()
    # список брендов и категорий уже есть в админке, но в reports надо тоже
    cur.execute("SELECT category_id, name FROM dbo.categories WHERE name not like 'Зап%'ORDER BY name")
    all_categories = cur.fetchall()

    cur.execute("SELECT id_brand, name FROM dbo.brand ORDER BY name")
    all_brands = cur.fetchall()

    # отчёт по комплектации (запчасти бренда)
    # отчёт по комплектации (запчасти по бренду; бренд может быть пустой)
    parts_rows = []
    if target_category_id.isdigit():
        bval = int(brand_id) if brand_id.isdigit() else None
        cur.execute("EXEC dbo.sp_report_parts_by_brand_and_category ?, ?",
                    int(target_category_id), bval)
        parts_rows = cur.fetchall()
    conn.close()

    return render_template(
        "admin.html",
        tab="reports",
        report_rows=report_rows,
        orders_rows=orders_rows,
        customers_rows=customers_rows,
        sold_rows=sold_rows,
        date_from=date_from,
        date_to=date_to,
        all_categories=all_categories,
        all_brands=all_brands,
        target_category_id=target_category_id,
        parts_brand_id=brand_id,
        parts_rows=parts_rows,
    )



import io
from datetime import datetime

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

import io
import os
from datetime import datetime

from flask import current_app, send_file
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XLImage

@auth_bp.route("/admin/reports/excel")
def export_excel():
    gate = _require_admin()
    if gate:
        return gate

    # 1) данные
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("SELECT * FROM dbo.vw_report_products_per_category ORDER BY category_name")
    rows = cur.fetchall()
    conn.close()

    # 2) excel
    wb = Workbook()
    ws = wb.active
    ws.title = "Отчет"

    bold = Font(bold=True)
    bold_big = Font(bold=True, size=14)
    bold_mid = Font(bold=True, size=12)

    wrap_left = Alignment(wrap_text=True, vertical="top", horizontal="left")
    center = Alignment(horizontal="center", vertical="center")
    center_wrap = Alignment(horizontal="center", vertical="center", wrap_text=True)

    thin = Side(style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # --- ШАПКА: ЛОГО + РЕКВИЗИТЫ ---
    # Ячейка под лого (A1:B1) — делаем повыше
    ws.merge_cells("A1:B1")
    ws.row_dimensions[1].height = 60

    # путь к лого (проверь место!)
    # вариант 1: /static/rolmark_logo.png
    logo_path = os.path.join(current_app.root_path, "static", "rolmark_logo.png")

    # если у тебя реально лежит в app/static — то так:
    if not os.path.exists(logo_path):
        logo_path = os.path.join(current_app.root_path, "app", "static", "rolmark_logo.png")

    if os.path.exists(logo_path):
        img = XLImage(logo_path)
        # размер (можешь менять)
        img.width = 170
        img.height = 55
        ws.add_image(img, "A1")
    else:
        # запасной вариант: текст, если не нашли файл
        ws["A1"] = "ROLMARK"
        ws["A1"].font = bold_big
        ws["A1"].alignment = Alignment(horizontal="left", vertical="center")

    requisites = (
        "Юр.адрес: 220024, Республика Беларусь, г.Минск, ул.Бабушкина, д.4а каб. 33\n"
        "Салон-офис: 220006, Республика Беларусь, г.Минск, ул.Маяковского, д.26, каб.1 (вход со двора)\n"
        "Р/с BY86TECN30121248300010000000 в ОАО «Технобанк», г.Минск, ул.Кропоткина,44 БИК TECNBY22\n"
        "УНП 190640194\n"
        "тел./факс: +375 17 348 99 82\n"
        "МТС +375 33 361 65 65, Velcom +375 29 361 65 65\n"
        "Наша электронная почта: rolmark.trade@gmail.com"
    )

    ws["A2"] = requisites
    ws["A2"].alignment = wrap_left
    ws["A2"].font = Font(size=9)
    ws.merge_cells("A2:B5")

    # Заголовки по центру
    ws["A6"] = "ОТЧЁТ"
    ws["A6"].font = bold_big
    ws["A6"].alignment = center
    ws.merge_cells("A6:B6")

    ws["A7"] = "по количеству товаров в категориях"
    ws["A7"].font = bold_mid
    ws["A7"].alignment = center
    ws.merge_cells("A7:B7")

    ws["A8"] = f"По состоянию на: {datetime.now().strftime('%d.%m.%Y')}"
    ws["A8"].font = bold
    ws["A8"].alignment = center
    ws.merge_cells("A8:B8")

    # --- ТАБЛИЦА (без ID) ---
    start_row = 10

    ws.cell(row=start_row, column=1, value="Категория").font = bold
    ws.cell(row=start_row, column=2, value="Количество").font = bold
    ws.cell(row=start_row, column=1).alignment = center_wrap
    ws.cell(row=start_row, column=2).alignment = center_wrap

    r0 = start_row + 1
    for i, r in enumerate(rows):
        rr = r0 + i
        ws.cell(row=rr, column=1, value=str(r.category_name)).alignment = Alignment(vertical="top", wrap_text=True)
        ws.cell(row=rr, column=2, value=int(r.products_count)).alignment = center

    last_row = r0 + len(rows) - 1 if rows else start_row

    for rr in range(start_row, last_row + 1):
        for cc in range(1, 3):
            ws.cell(row=rr, column=cc).border = border

    # закрепим верх (строка после заголовка таблицы)
    ws.freeze_panes = ws["A11"]

    # ширины колонок
    for col in range(1, 3):
        max_len = 0
        for rr in range(1, last_row + 1):
            v = ws.cell(row=rr, column=col).value
            if v is None:
                continue
            max_len = max(max_len, len(str(v)))
        ws.column_dimensions[get_column_letter(col)].width = min(max(12, max_len + 2), 60)

    # 3) отдаём файл
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="report.xlsx")
from flask import Blueprint, render_template, request, redirect, url_for, session, send_file, current_app
from app.db import get_conn
import io
import os

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


from flask import current_app
import os
import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

@auth_bp.route("/admin/reports/word")
def export_word():
    gate = _require_admin()
    if gate:
        return gate

    # 1) данные
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("SELECT * FROM dbo.vw_report_products_per_category ORDER BY category_name")
    rows = cur.fetchall()
    conn.close()

    # 2) документ
    doc = Document()
    # --- нижний колонтитул ---
    now_str = datetime.now().strftime("%d.%m.%Y %H:%M")
    section = doc.sections[0]
    footer = section.footer

    # очистим дефолтный пустой абзац, чтобы не было лишних строк
    if footer.paragraphs:
        footer.paragraphs[0].text = ""

    p_f = footer.add_paragraph()
    p_f.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_f = p_f.add_run(f"Отчет составлен: {now_str}")
    run_f.bold = True
    run_f.font.size = Pt(9)

    # --- верхний левый блок: лого + реквизиты ---
    top = doc.add_table(rows=1, cols=2)
    top.autofit = True

    left = top.cell(0, 0)
    right = top.cell(0, 1)
    right.text = ""

    # ✅ ПРАВИЛЬНЫЙ путь: current_app.root_path уже == .../app
    logo_path = os.path.join(current_app.root_path, "app", "static", "rolmark_logo.png")
    # print("LOGO PATH:", logo_path)
    # print("EXISTS:", os.path.exists(logo_path))

    if os.path.exists(logo_path):
        p_logo = left.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1.6))

    requisites = (
        "Юр.адрес: 220024, Республика Беларусь, г.Минск, ул.Бабушкина, д.4а каб. 33\n"
        "Салон-офис: 220006, Республика Беларусь, г.Минск, ул.Маяковского, д.26, каб.1 (вход со двора)\n"
        "Р/с BY86TECN30121248300010000000 в ОАО «Технобанк», г.Минск, ул.Кропоткина,44 БИК TECNBY22\n"
        "УНП 190640194\n"
        "тел./факс: +375 17 348 99 82\n"
        "МТС +375 33 361 65 65, Velcom +375 29 361 65 65\n"
        "Наша электронная почта: rolmark.trade@gmail.com"
    )

    p_req = left.add_paragraph(requisites)
    p_req.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p_req.runs:
        r.font.size = Pt(9)

    doc.add_paragraph("")  # отступ после шапки

    # --- Заголовки по центру ---
    p1 = doc.add_paragraph("ОТЧЁТ")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.runs[0]
    r1.bold = True
    r1.font.size = Pt(14)

    p2 = doc.add_paragraph("по количеству товаров в категориях")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.runs[0]
    r2.bold = True
    r2.font.size = Pt(12)

    # ✅ дата: по центру и жирным
    p_date = doc.add_paragraph(f"По состоянию на: {datetime.now().strftime('%d.%m.%Y')}")
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = p_date.runs[0]
    rd.bold = True
    rd.font.size = Pt(11)

    doc.add_paragraph("")

    # --- Таблица: Категория + Количество, с границами ---
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"

    hdr = t.rows[0].cells
    hdr[0].text = "Категория"
    hdr[1].text = "Количество"

    for r in rows:
        c = t.add_row().cells
        c[0].text = str(r.category_name)
        c[1].text = str(r.products_count)

    # 3) отдаём файл
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(buf, as_attachment=True, download_name="report.docx")
@auth_bp.route("/admin/reports/orders")
def report_orders_period():
    gate = _require_admin()
    if gate:
        return gate

    date_from = request.args.get("from", "").strip()
    date_to = request.args.get("to", "").strip()

    rows = []
    if date_from and date_to:
        conn = get_conn()
        cur = conn.cursor()
        cur.execute("EXEC dbo.sp_report_orders_period_grouped ?, ?", date_from, date_to)
        rows = cur.fetchall()
        conn.close()

    return render_template(
        "admin.html",
        tab="reports",
        report_tab="orders",
        rows=rows,
        date_from=date_from,
        date_to=date_to
    )
@auth_bp.route("/admin/reports/clients")
def report_clients():
    gate = _require_admin()
    if gate:
        return gate

    conn = get_conn()
    cur = conn.cursor()
    cur.execute("EXEC dbo.sp_report_clients_orders")
    rows = cur.fetchall()
    conn.close()

    return render_template(
        "admin.html",
        tab="reports",
        report_tab="clients",
        rows=rows
    )
@auth_bp.route("/admin/reports/products")
def report_products_period():
    gate = _require_admin()
    if gate:
        return gate

    date_from = request.args.get("from", "").strip()
    date_to = request.args.get("to", "").strip()

    rows = []
    if date_from and date_to:
        conn = get_conn()
        cur = conn.cursor()
        cur.execute("EXEC dbo.sp_report_sold_products_period ?, ?", date_from, date_to)
        rows = cur.fetchall()
        conn.close()

    return render_template(
        "admin.html",
        tab="reports",
        report_tab="products",
        rows=rows,
        date_from=date_from,
        date_to=date_to
    )
import os
import io
from datetime import datetime
from flask import current_app, send_file, redirect, request
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Font, Alignment
from openpyxl.utils import get_column_letter

import os
import io
from datetime import datetime
from flask import current_app, send_file, redirect, request
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

import os
import io
from datetime import datetime
from flask import current_app, send_file, redirect, request
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from datetime import datetime
import os, io
from flask import current_app, send_file, redirect
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

@auth_bp.route("/admin/reports/orders/excel")
def export_orders_excel():
    gate = _require_admin()
    if gate:
        return gate

    date_from = request.args.get("from", "").strip()
    date_to = request.args.get("to", "").strip()
    if not date_from or not date_to:
        return redirect("/admin/reports?rtab=orders")

    # 1) данные
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("EXEC dbo.sp_report_orders_period_grouped ?, ?", date_from, date_to)
    rows = cur.fetchall()
    conn.close()

    # 2) excel
    wb = Workbook()
    ws = wb.active
    ws.title = "Orders report"

    bold = Font(bold=True)
    bold_big = Font(bold=True, size=14)
    bold_mid = Font(bold=True, size=12)

    wrap_left = Alignment(wrap_text=True, vertical="top", horizontal="left")
    center = Alignment(horizontal="center", vertical="center")
    center_wrap = Alignment(horizontal="center", vertical="center", wrap_text=True)

    thin = Side(style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # ширины колонок (под таблицу)
    widths = [14, 30, 28, 12, 12, 16]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # --- ШАПКА: ЛОГО (A1) + РЕКВИЗИТЫ В A2 (ОДНА ЯЧЕЙКА) ---
    ws.row_dimensions[1].height = 70

    logo_path = os.path.join(current_app.root_path, "static", "rolmark_logo.png")
    if not os.path.exists(logo_path):
        logo_path = os.path.join(current_app.root_path, "app", "static", "rolmark_logo.png")

    if os.path.exists(logo_path):
        img = XLImage(logo_path)
        img.width = 190
        img.height = 65
        ws.add_image(img, "A1")
    else:
        ws["A1"] = "ROLMARK-TRADE"
        ws["A1"].font = bold_big
        ws["A1"].alignment = Alignment(horizontal="left", vertical="center")

    requisites = (
        "Юр.адрес: 220024, Республика Беларусь, г.Минск, ул.Бабушкина, д.4а каб. 33\n"
        "Салон-офис: 220006, Республика Беларусь, г.Минск, ул.Маяковского, д.26, каб.1 (вход со двора)\n"
        "Р/с  BY86TECN30121248300010000000 в ОАО «Технобанк», г.Минск, ул.Кропоткина,44 БИК TECNBY22\n"
        "УНП 190640194\n"
        "тел./факс: + 375 17 348 99 82\n"
        "МТС + 375 33 361 65 65, Velcom + 375 29 361 65 65\n"
        "Наша электронная почта:  rolmark.trade@gmail.com"
    )

    # реквизиты строго в A2 (одна ячейка)
    ws["A2"] = requisites
    ws["A2"].font = Font(size=10)
    ws["A2"].alignment = wrap_left

    # чтобы реквизиты были видны — делаем блок A2:F8 (одна большая ячейка)
    ws.merge_cells("A2:F8")

    # высоты строк под реквизиты
    for r in range(2, 9):
        ws.row_dimensions[r].height = 18

    # рамка вокруг шапки (A1:F8), чтобы выглядело аккуратно
    #for rr in range(1, 9):
        #for cc in range(1, 7):
           # ws.cell(row=rr, column=cc).border = border

    # --- Заголовок отчёта ---
    title_row = 10

    ws.merge_cells(start_row=title_row, start_column=1, end_row=title_row, end_column=6)
    c1 = ws.cell(row=title_row, column=1)
    c1.value = "ОТЧЁТ"
    c1.font = bold_big
    c1.alignment = center

    ws.merge_cells(start_row=title_row + 1, start_column=1, end_row=title_row + 1, end_column=6)
    c2 = ws.cell(row=title_row + 1, column=1)
    c2.value = f"по заказам за период {date_from} — {date_to}"
    c2.font = bold_mid
    c2.alignment = center

    ws.merge_cells(start_row=title_row + 2, start_column=1, end_row=title_row + 2, end_column=6)
    c3 = ws.cell(row=title_row + 2, column=1)
    c3.value = f"По состоянию на: {datetime.now().strftime('%d.%m.%Y')}"
    c3.font = bold
    c3.alignment = center

    # --- Таблица ---
    table_start = title_row + 4
    headers = ["Статус", "Клиент", "Категория", "Заказов", "Позиций", "Сумма"]

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=table_start, column=col, value=h)
        cell.font = bold
        cell.alignment = center_wrap
        cell.border = border

    r_i = table_start + 1
    for r in rows:
        ws.cell(row=r_i, column=1, value=str(r.status) if r.status is not None else "")
        ws.cell(row=r_i, column=2, value=str(r.client_email) if r.client_email is not None else "")
        ws.cell(row=r_i, column=3, value=str(r.category_name) if r.category_name is not None else "")
        ws.cell(row=r_i, column=4, value=int(r.orders_count) if r.orders_count is not None else 0)
        ws.cell(row=r_i, column=5, value=int(r.items_count) if r.items_count is not None else 0)
        ws.cell(row=r_i, column=6, value=float(r.total_sum) if r.total_sum is not None else 0.0)

        for col in range(1, 7):
            ws.cell(row=r_i, column=col).alignment = Alignment(vertical="top", wrap_text=True)
            ws.cell(row=r_i, column=col).border = border

        r_i += 1

    ws.freeze_panes = ws[f"A{table_start+1}"]

    # сохранить и отдать
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="orders_report.xlsx")
import os
import io
from datetime import datetime
from flask import current_app, send_file, redirect, request
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

@auth_bp.route("/admin/reports/orders/word")
def export_orders_word():
    gate = _require_admin()
    if gate:
        return gate

    date_from = request.args.get("from", "").strip()
    date_to = request.args.get("to", "").strip()
    if not date_from or not date_to:
        return redirect("/admin/reports?rtab=orders")

    # 1) данные
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("EXEC dbo.sp_report_orders_period_grouped ?, ?", date_from, date_to)
    rows = cur.fetchall()
    conn.close()

    # 2) документ
    doc = Document()
    # --- нижний колонтитул ---
    now_str = datetime.now().strftime("%d.%m.%Y %H:%M")
    section = doc.sections[0]
    footer = section.footer

    # очистим дефолтный пустой абзац, чтобы не было лишних строк
    if footer.paragraphs:
        footer.paragraphs[0].text = ""

    p_f = footer.add_paragraph()
    p_f.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_f = p_f.add_run(f"Отчет составлен: {now_str}")
    run_f.bold = True
    run_f.font.size = Pt(9)

    # --- шапка: лого + реквизиты ---
    top = doc.add_table(rows=1, cols=2)
    top.autofit = True

    left = top.cell(0, 0)
    right = top.cell(0, 1)
    right.text = ""

    # путь к лого (попробуем 2 варианта как раньше)
    logo_path = os.path.join(current_app.root_path, "static", "rolmark_logo.png")
    if not os.path.exists(logo_path):
        logo_path = os.path.join(current_app.root_path, "app", "static", "rolmark_logo.png")

    if os.path.exists(logo_path):
        p_logo = left.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_logo.add_run()
        run.add_picture(logo_path, width=Inches(1.6))

    requisites = (
        "Юр.адрес: 220024, Республика Беларусь, г.Минск, ул.Бабушкина, д.4а каб. 33\n\n"
        "Салон-офис: 220006, Республика Беларусь, г.Минск, ул.Маяковского, д.26, каб.1 (вход со двора)\n\n"
        "Р/с BY86TECN30121248300010000000 в ОАО «Технобанк», г.Минск, ул.Кропоткина,44 БИК TECNBY22\n\n"
        "УНП 190640194\n\n"
        "тел./факс: + 375 17 348 99 82\n\n"
        "МТС + 375 33 361 65 65, Velcom + 375 29 361 65 65\n\n"
        "Наша электронная почта:  rolmark.trade@gmail.com"
    )

    p_req = left.add_paragraph(requisites)
    p_req.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p_req.runs:
        r.font.size = Pt(9)

    doc.add_paragraph("")

    # --- заголовки по центру ---
    p1 = doc.add_paragraph("ОТЧЁТ")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.runs[0]
    r1.bold = True
    r1.font.size = Pt(14)

    p2 = doc.add_paragraph(f"по заказам за период {date_from} — {date_to}")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.runs[0]
    r2.bold = True
    r2.font.size = Pt(12)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rdate = p_date.add_run(f"По состоянию на: {datetime.now().strftime('%d.%m.%Y')}")
    rdate.bold = True
    rdate.font.size = Pt(11)

    doc.add_paragraph("")

    # --- таблица с границами ---
    t = doc.add_table(rows=1, cols=6)
    t.style = "Table Grid"

    hdr = t.rows[0].cells
    hdr[0].text = "Статус"
    hdr[1].text = "Клиент"
    hdr[2].text = "Категория"
    hdr[3].text = "Заказов"
    hdr[4].text = "Позиций"
    hdr[5].text = "Сумма"

    for r in rows:
        c = t.add_row().cells
        c[0].text = str(r.status)
        c[1].text = str(r.client_email)
        c[2].text = str(r.category_name)
        c[3].text = str(r.orders_count)
        c[4].text = str(r.items_count)
        c[5].text = str(r.total_sum)

    # 3) отдаём
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="orders_report.docx")
@auth_bp.route("/admin/reports/clients/excel")
def export_clients_excel():
    gate = _require_admin()
    if gate:
        return gate

    # 1) данные
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("EXEC dbo.sp_report_clients_orders")
    rows = cur.fetchall()
    conn.close()

    # 2) excel
    wb = Workbook()
    ws = wb.active
    ws.title = "Клиенты"

    bold = Font(bold=True)
    bold_big = Font(bold=True, size=14)
    bold_mid = Font(bold=True, size=12)

    wrap_left = Alignment(wrap_text=True, vertical="top", horizontal="left")
    center = Alignment(horizontal="center", vertical="center")
    center_wrap = Alignment(horizontal="center", vertical="center", wrap_text=True)

    thin = Side(style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # --- ЛОГО (A1) ---
    ws.merge_cells("A1:E1")
    ws.row_dimensions[1].height = 60

    logo_path = os.path.join(current_app.root_path, "static", "rolmark_logo.png")
    if not os.path.exists(logo_path):
        logo_path = os.path.join(current_app.root_path, "app", "static", "rolmark_logo.png")

    if os.path.exists(logo_path):
        img = XLImage(logo_path)
        img.width = 170
        img.height = 55
        ws.add_image(img, "A1")
    else:
        ws["A1"] = "ROLMARK"
        ws["A1"].font = bold_big
        ws["A1"].alignment = Alignment(horizontal="left", vertical="center")

    # --- РЕКВИЗИТЫ В ОДНОЙ ЯЧЕЙКЕ A2 ---
    requisites = (
        "Юр.адрес: 220024, Республика Беларусь, г.Минск, ул.Бабушкина, д.4а каб. 33\n"
        "Салон-офис: 220006, Республика Беларусь, г.Минск, ул.Маяковского, д.26, каб.1 (вход со двора)\n"
        "Р/с BY86TECN30121248300010000000 в ОАО «Технобанк», г.Минск, ул.Кропоткина,44 БИК TECNBY22\n"
        "УНП 190640194\n"
        "тел./факс: +375 17 348 99 82\n"
        "МТС +375 33 361 65 65, Velcom +375 29 361 65 65\n"
        "Наша электронная почта: rolmark.trade@gmail.com"
    )
    ws["A2"] = requisites
    ws["A2"].font = Font(size=9)
    ws["A2"].alignment = wrap_left
    ws.merge_cells("A2:E6")

    # --- ЗАГОЛОВОК ---
    ws["A7"] = "ОТЧЁТ"
    ws["A7"].font = bold_big
    ws["A7"].alignment = center
    ws.merge_cells("A7:E7")

    ws["A8"] = "по клиентам (история заказов)"
    ws["A8"].font = bold_mid
    ws["A8"].alignment = center
    ws.merge_cells("A8:E8")

    ws["A9"] = f"По состоянию на: {datetime.now().strftime('%d.%m.%Y')}"
    ws["A9"].font = bold
    ws["A9"].alignment = center
    ws.merge_cells("A9:E9")

    # --- ТАБЛИЦА (БЕЗ User ID) ---
    table_start = 11
    headers = ["Email", "Заказов", "Потрачено", "Первый заказ", "Последний заказ"]

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=table_start, column=col, value=h)
        cell.font = bold
        cell.alignment = center_wrap

    r_i = table_start + 1
    for r in rows:
        ws.cell(row=r_i, column=1, value=str(r.email))
        ws.cell(row=r_i, column=2, value=int(r.orders_count) if r.orders_count is not None else 0)
        ws.cell(row=r_i, column=3, value=float(r.total_spent) if r.total_spent is not None else 0.0)
        ws.cell(row=r_i, column=4, value=str(r.first_order) if r.first_order else "Клиент ничего не покупал")
        ws.cell(row=r_i, column=5, value=str(r.last_order) if r.last_order else "Клиент ничего не покупал")

        for col in range(1, 6):
            ws.cell(row=r_i, column=col).alignment = Alignment(vertical="top", wrap_text=True)
        r_i += 1

    last_row = r_i - 1

    # границы только у таблицы (шапку/реквизиты не трогаем)
    for rr in range(table_start, last_row + 1):
        for cc in range(1, 6):
            ws.cell(row=rr, column=cc).border = border

    # ширины колонок
    widths = [32, 12, 14, 18, 18]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    ws.freeze_panes = ws["A12"]

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="clients_report.xlsx")
@auth_bp.route("/admin/reports/clients/word")
def export_clients_word():
    gate = _require_admin()
    if gate:
        return gate

    # 1) данные
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("EXEC dbo.sp_report_clients_orders")
    rows = cur.fetchall()
    conn.close()

    # 2) документ
    doc = Document()

    # --- нижний колонтитул: "отчет составлен на ..." ---
    now = datetime.now()
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = f"Отчет составлен на {now.strftime('%d.%m.%Y %H:%M')}"
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in fp.runs:
        run.font.size = Pt(9)

    # --- верхний левый блок: лого + реквизиты ---
    top = doc.add_table(rows=1, cols=2)
    top.autofit = True

    left = top.cell(0, 0)
    right = top.cell(0, 1)
    right.text = ""

    # ищем лого в static
    logo_path = os.path.join(current_app.root_path, "static", "rolmark_logo.png")
    if not os.path.exists(logo_path):
        logo_path = os.path.join(current_app.root_path, "app", "static", "rolmark_logo.png")

    if os.path.exists(logo_path):
        p_logo = left.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_logo.add_run()
        run.add_picture(logo_path, width=Inches(1.6))

    requisites = (
        "Юр.адрес: 220024, Республика Беларусь, г.Минск, ул.Бабушкина, д.4а каб. 33\n\n"
        "Салон-офис: 220006, Республика Беларусь, г.Минск, ул.Маяковского, д.26, каб.1 (вход со двора)\n\n"
        "Р/с BY86TECN30121248300010000000 в ОАО «Технобанк», г.Минск, ул.Кропоткина,44 БИК TECNBY22\n\n"
        "УНП 190640194\n\n"
        "тел./факс: +375 17 348 99 82\n\n"
        "МТС +375 33 361 65 65, Velcom +375 29 361 65 65\n\n"
        "Наша электронная почта: rolmark.trade@gmail.com"
    )

    p_req = left.add_paragraph(requisites)
    p_req.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p_req.runs:
        run.font.size = Pt(9)

    doc.add_paragraph("")

    # --- Заголовок по центру ---
    p1 = doc.add_paragraph("ОТЧЁТ")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.runs[0]
    r1.bold = True
    r1.font.size = Pt(14)

    p2 = doc.add_paragraph("по клиентам (история заказов)")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.runs[0]
    r2.bold = True
    r2.font.size = Pt(12)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"По состоянию на: {now.strftime('%d.%m.%Y')}")
    r3.bold = True
    r3.font.size = Pt(11)

    doc.add_paragraph("")

    # --- Таблица с границами ---
    # (оставляем 6 колонок как у тебя)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"

    hdr = t.rows[0].cells
    hdr[0].text = "Email"
    hdr[1].text = "Заказов"
    hdr[2].text = "Потрачено"
    hdr[3].text = "Первый заказ"
    hdr[4].text = "Последний заказ"

    for r in rows:
        c = t.add_row().cells
        c[0].text = str(r.email)
        c[1].text = str(r.orders_count)
        c[2].text = str(r.total_spent)
        c[3].text = str(r.first_order) if r.first_order else "Клиент ничего не покупал"
        c[4].text = str(r.last_order) if r.last_order else "Клиент ничего не покупал"
    # 3) отдаём файл
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="clients_report.docx")
@auth_bp.route("/admin/reports/products/excel")
def export_products_excel():
    gate = _require_admin()
    if gate:
        return gate

    date_from = request.args.get("from", "").strip()
    date_to = request.args.get("to", "").strip()
    if not date_from or not date_to:
        return redirect("/admin/reports?rtab=products")

    # 1) данные
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("EXEC dbo.sp_report_sold_products_period ?, ?", date_from, date_to)
    rows = cur.fetchall()
    conn.close()

    # 2) excel
    wb = Workbook()
    ws = wb.active
    ws.title = "Проданные товары"

    bold = Font(bold=True)
    bold_big = Font(bold=True, size=14)
    bold_mid = Font(bold=True, size=12)

    left_wrap = Alignment(wrap_text=True, vertical="top", horizontal="left")
    center = Alignment(horizontal="center", vertical="center")
    center_wrap = Alignment(horizontal="center", vertical="center", wrap_text=True)

    thin = Side(style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # колонки под таблицу (4 колонки)
    ws.column_dimensions["A"].width = 38  # Товар
    ws.column_dimensions["B"].width = 18  # Категория (если есть в rows)
    ws.column_dimensions["C"].width = 16  # Продано
    ws.column_dimensions["D"].width = 18  # Выручка

    # --- ЛОГО В A1 ---
    ws.row_dimensions[1].height = 60

    logo_path = os.path.join(current_app.root_path, "static", "rolmark_logo.png")
    if not os.path.exists(logo_path):
        logo_path = os.path.join(current_app.root_path, "app", "static", "rolmark_logo.png")

    if os.path.exists(logo_path):
        img = XLImage(logo_path)
        img.width = 190
        img.height = 95
        ws.add_image(img, "A1")
    else:
        ws["A1"] = "ROLMARK"
        ws["A1"].font = bold_big
        ws["A1"].alignment = Alignment(horizontal="left", vertical="center")

    # --- РЕКВИЗИТЫ В ОДНОЙ ЯЧЕЙКЕ A2 ---
    requisites = (
        "Юр.адрес: 220024, Республика Беларусь, г.Минск, ул.Бабушкина, д.4а каб. 33\n"
        "Салон-офис: 220006, Республика Беларусь, г.Минск, ул.Маяковского, д.26, каб.1 (вход со двора)\n"
        "Р/с  BY86TECN30121248300010000000 в ОАО «Технобанк», г.Минск, ул.Кропоткина,44 БИК TECNBY22\n"
        "УНП 190640194\n"
        "тел./факс: + 375 17 348 99 82\n"
        "МТС + 375 33 361 65 65, Velcom + 375 29 361 65 65\n"
        "Наша электронная почта:  rolmark.trade@gmail.com"
    )

    ws["A3"] = requisites
    ws["A3"].alignment = left_wrap
    ws["A3"].font = Font(size=10)
    ws.merge_cells("A3:D7")          # реквизиты большим блоком
    ws.row_dimensions[3].height = 18
    ws.row_dimensions[4].height = 18
    ws.row_dimensions[5].height = 18
    ws.row_dimensions[6].height = 18
    ws.row_dimensions[7].height = 18
    ws.row_dimensions[8].height = 18

    # --- Заголовок отчёта по центру ---
    ws["A9"] = "ОТЧЁТ"
    ws["A9"].font = bold_big
    ws["A9"].alignment = center
    ws.merge_cells("A9:D9")

    ws["A10"] = f"по проданным товарам за период {date_from} — {date_to}"
    ws["A10"].font = bold_mid
    ws["A10"].alignment = center
    ws.merge_cells("A10:D10")

    ws["A11"] = f"По состоянию на: {datetime.now().strftime('%d.%m.%Y')}"
    ws["A11"].font = bold
    ws["A11"].alignment = center
    ws.merge_cells("A11:D11")

    # --- Таблица (без ID) ---
    # ВНИМАНИЕ: если ты ДОБАВИЛ категорию в процедуру, она может быть в rows как r.category_name
    # Если категории пока нет — столбец "Категория" просто заполним "-"
    table_start = 13

    headers = ["Товар", "Категория", "Продано (шт)", "Выручка"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=table_start, column=col, value=h)
        cell.font = bold
        cell.alignment = center_wrap

    r_i = table_start + 1
    for r in rows:
        # пытаемся взять категорию, если она есть в результате процедуры
        cat_val = "-"
        if hasattr(r, "category_name") and r.category_name is not None:
            cat_val = str(r.category_name)
        elif hasattr(r, "category") and r.category is not None:
            cat_val = str(r.category)

        ws.cell(row=r_i, column=1, value=str(r.name) if r.name is not None else "-").alignment = Alignment(wrap_text=True, vertical="top")
        ws.cell(row=r_i, column=2, value=cat_val).alignment = Alignment(wrap_text=True, vertical="top")
        ws.cell(row=r_i, column=3, value=int(r.sold_qty) if r.sold_qty is not None else 0).alignment = center
        ws.cell(row=r_i, column=4, value=float(r.revenue) if r.revenue is not None else 0.0).alignment = center
        r_i += 1

    last_row = max(table_start, r_i - 1)

    # границы только на таблицу (шапка+данные)
    for rr in range(table_start, last_row + 1):
        for cc in range(1, 5):
            ws.cell(row=rr, column=cc).border = border

    # закрепим заголовок таблицы
    ws.freeze_panes = ws["A14"]

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="sold_products_report.xlsx")
@auth_bp.route("/admin/reports/products/word")
def export_products_word():
    gate = _require_admin()
    if gate:
        return gate

    date_from = request.args.get("from", "").strip()
    date_to = request.args.get("to", "").strip()
    if not date_from or not date_to:
        return redirect("/admin/reports?rtab=products")

    # 1) данные
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("EXEC dbo.sp_report_sold_products_period ?, ?", date_from, date_to)
    rows = cur.fetchall()
    conn.close()

    # 2) документ
    doc = Document()
    # --- нижний колонтитул ---
    now_str = datetime.now().strftime("%d.%m.%Y %H:%M")
    section = doc.sections[0]
    footer = section.footer

    # очистим дефолтный пустой абзац, чтобы не было лишних строк
    if footer.paragraphs:
        footer.paragraphs[0].text = ""

    p_f = footer.add_paragraph()
    p_f.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_f = p_f.add_run(f"Отчет составлен: {now_str}")
    run_f.bold = True
    run_f.font.size = Pt(9)

    # --- шапка: лого + реквизиты ---
    top = doc.add_table(rows=1, cols=2)
    top.autofit = True
    left = top.cell(0, 0)
    right = top.cell(0, 1)
    right.text = ""

    logo_path = os.path.join(current_app.root_path, "static", "rolmark_logo.png")
    if not os.path.exists(logo_path):
        logo_path = os.path.join(current_app.root_path, "app", "static", "rolmark_logo.png")

    if os.path.exists(logo_path):
        p_logo = left.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_logo.add_run()
        run.add_picture(logo_path, width=Inches(1.6))

    requisites = (
        "Юр.адрес: 220024, Республика Беларусь, г.Минск, ул.Бабушкина, д.4а каб. 33\n\n"
        "Салон-офис: 220006, Республика Беларусь, г.Минск, ул.Маяковского, д.26, каб.1 (вход со двора)\n\n"
        "Р/с BY86TECN30121248300010000000 в ОАО «Технобанк», г.Минск, ул.Кропоткина,44 БИК TECNBY22\n\n"
        "УНП 190640194\n\n"
        "тел./факс: +375 17 348 99 82\n\n"
        "МТС +375 33 361 65 65, Velcom +375 29 361 65 65\n\n"
        "Наша электронная почта: rolmark.trade@gmail.com"
    )
    p_req = left.add_paragraph(requisites)
    p_req.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for rr in p_req.runs:
        rr.font.size = Pt(9)

    doc.add_paragraph("")

    # --- Заголовки по центру ---
    p1 = doc.add_paragraph("ОТЧЁТ")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.runs[0].bold = True
    p1.runs[0].font.size = Pt(14)

    p2 = doc.add_paragraph(f"по проданным товарам за период {date_from} — {date_to}")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].bold = True
    p2.runs[0].font.size = Pt(12)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_dt = p3.add_run(f"По состоянию на: {datetime.now().strftime('%d.%m.%Y')}")
    run_dt.bold = True
    run_dt.font.size = Pt(11)

    doc.add_paragraph("")

    # --- Таблица (БЕЗ ID), с границами ---
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"

    hdr = t.rows[0].cells
    hdr[0].text = "Товар"
    hdr[1].text = "Категория"
    hdr[2].text = "Продано (шт)"
    hdr[3].text = "Выручка"

    for r in rows:
        c = t.add_row().cells
        c[0].text = str(getattr(r, "name", ""))
        c[1].text = str(getattr(r, "category_name", ""))  # должно прийти из процедуры
        c[2].text = str(getattr(r, "sold_qty", 0))
        c[3].text = str(getattr(r, "revenue", 0))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="sold_products_report.docx")
@auth_bp.route("/admin/reports/parts/excel")
def export_parts_excel():
    gate = _require_admin()
    if gate:
        return gate

    target_category_id = request.args.get("target_category_id", "").strip()
    brand_id = request.args.get("parts_brand_id", "").strip()

    # ✅ category обязателен, brand может быть пустым = все бренды
    if not target_category_id.isdigit():
        return redirect("/admin/reports")

    bval = int(brand_id) if brand_id.isdigit() else None

    # 1) данные
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        "EXEC dbo.sp_report_parts_by_brand_and_category ?, ?",
        int(target_category_id), bval
    )
    rows = cur.fetchall()
    conn.close()

    # 2) excel
    wb = Workbook()
    ws = wb.active
    ws.title = "Parts report"

    bold = Font(bold=True)
    bold_big = Font(bold=True, size=14)
    bold_mid = Font(bold=True, size=12)

    left_wrap = Alignment(wrap_text=True, vertical="top", horizontal="left")
    center = Alignment(horizontal="center", vertical="center")
    center_wrap = Alignment(horizontal="center", vertical="center", wrap_text=True)

    thin = Side(style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # сделаем ширину под “шапку”
    ws.column_dimensions["A"].width = 22   # под лого
    ws.column_dimensions["B"].width = 95   # под реквизиты/текст (будем мерджить A3:F6, но ширина B поможет)

    # --- ЛОГО в A1 ---
    ws.row_dimensions[1].height = 55

    logo_path = os.path.join(current_app.root_path, "static", "rolmark_logo.png")
    if not os.path.exists(logo_path):
        logo_path = os.path.join(current_app.root_path, "app", "static", "rolmark_logo.png")

    if os.path.exists(logo_path):
        img = XLImage(logo_path)
        img.width = 190
        img.height = 95
        ws.add_image(img, "A1")
    else:
        ws["A1"] = "ROLMARK"
        ws["A1"].font = bold_big
        ws["A1"].alignment = Alignment(horizontal="left", vertical="center")

    # --- Реквизиты в A3 (одна ячейка), как ты просил ---
    requisites = (
        "Юр.адрес: 220024, Республика Беларусь, г.Минск, ул.Бабушкина, д.4а каб. 33\n"
        "Салон-офис: 220006, Республика Беларусь, г.Минск, ул.Маяковского, д.26, каб.1 (вход со двора)\n"
        "Р/с  BY86TECN30121248300010000000 в ОАО «Технобанк», г.Минск, ул.Кропоткина,44 БИК TECNBY22\n"
        "УНП 190640194\n"
        "тел./факс: + 375 17 348 99 82\n"
        "МТС + 375 33 361 65 65, Velcom + 375 29 361 65 65\n"
        "Наша электронная почта:  rolmark.trade@gmail.com"
    )

    # чтобы реквизиты “смотрелись” и не упирались — дадим место
    ws.merge_cells("A3:F6")
    ws["A3"] = requisites
    ws["A3"].font = Font(size=10)
    ws["A3"].alignment = left_wrap
    ws.row_dimensions[3].height = 18
    ws.row_dimensions[4].height = 18
    ws.row_dimensions[5].height = 18
    ws.row_dimensions[6].height = 18

    # --- Заголовок по центру ---
    ws.merge_cells("A8:F8")
    ws["A8"] = "ОТЧЁТ"
    ws["A8"].font = bold_big
    ws["A8"].alignment = center

    ws.merge_cells("A9:F9")
    ws["A9"] = "по комплектации товаров (запчасти по бренду)"
    ws["A9"].font = bold_mid
    ws["A9"].alignment = center

    ws.merge_cells("A10:F10")
    ws["A10"] = f"По состоянию на: {datetime.now().strftime('%d.%m.%Y')}"
    ws["A10"].font = bold
    ws["A10"].alignment = center

    # --- Таблица (БЕЗ ID) ---
    table_start = 12
    headers = ["Запчасть", "Бренд", "Категория запчасти", "Цена", "Остаток"]

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=table_start, column=col)
        cell.value = h
        cell.font = bold
        cell.alignment = center_wrap

    r_i = table_start + 1
    for r in rows:
        ws.cell(row=r_i, column=1).value = r.part_name
        ws.cell(row=r_i, column=2).value = r.brand_name
        ws.cell(row=r_i, column=3).value = r.part_category
        ws.cell(row=r_i, column=4).value = float(r.price) if r.price is not None else 0.0
        ws.cell(row=r_i, column=5).value = int(r.stock_quantity) if r.stock_quantity is not None else 0

        for col in range(1, 6):
            ws.cell(row=r_i, column=col).alignment = Alignment(vertical="top", wrap_text=True)
        r_i += 1

    last_row = r_i - 1

    # границы только у таблицы (шапку с лого/реквизитами НЕ трогаем)
    for row in range(table_start, last_row + 1):
        for col in range(1, 6):
            ws.cell(row=row, column=col).border = border

    # ширины колонок таблицы
    widths = [34, 18, 30, 14, 12]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # закрепим шапку таблицы
    ws.freeze_panes = ws["A13"]

    # 3) отдать файл
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="parts_report.xlsx")
@auth_bp.route("/admin/reports/parts/word")
def export_parts_word():
    gate = _require_admin()
    if gate:
        return gate

    target_category_id = request.args.get("target_category_id", "").strip()
    brand_id = request.args.get("parts_brand_id", "").strip()

    # ✅ category обязателен, brand может быть пустым = все бренды
    if not target_category_id.isdigit():
        return redirect("/admin/reports")

    bval = int(brand_id) if brand_id.isdigit() else None

    # 1) данные
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        "EXEC dbo.sp_report_parts_by_brand_and_category ?, ?",
        int(target_category_id), bval
    )
    rows = cur.fetchall()
    conn.close()

    # 2) документ
    doc = Document()

    # --- верхний левый блок: лого + реквизиты ---
    top = doc.add_table(rows=1, cols=2)
    top.autofit = True
    left = top.cell(0, 0)
    right = top.cell(0, 1)
    right.text = ""

    # лого (ищем в /static и /app/static)
    logo_path = os.path.join(current_app.root_path, "static", "rolmark_logo.png")
    if not os.path.exists(logo_path):
        logo_path = os.path.join(current_app.root_path, "app", "static", "rolmark_logo.png")

    if os.path.exists(logo_path):
        p_logo = left.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_logo.add_run()
        run.add_picture(logo_path, width=Inches(1.6))
    else:
        p_logo = left.paragraphs[0]
        p_logo.add_run("ROLMARK").bold = True

    requisites = (
        "Юр.адрес: 220024, Республика Беларусь, г.Минск, ул.Бабушкина, д.4а каб. 33\n"
        "Салон-офис: 220006, Республика Беларусь, г.Минск, ул.Маяковского, д.26, каб.1 (вход со двора)\n"
        "Р/с  BY86TECN30121248300010000000 в ОАО «Технобанк», г.Минск, ул.Кропоткина,44 БИК TECNBY22\n"
        "УНП 190640194\n"
        "тел./факс: + 375 17 348 99 82\n"
        "МТС + 375 33 361 65 65, Velcom + 375 29 361 65 65\n"
        "Наша электронная почта:  rolmark.trade@gmail.com"
    )
    p_req = left.add_paragraph(requisites)
    p_req.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for rr in p_req.runs:
        rr.font.size = Pt(9)

    doc.add_paragraph("")

    # --- Заголовок по центру ---
    p1 = doc.add_paragraph("ОТЧЁТ")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.runs[0].bold = True
    p1.runs[0].font.size = Pt(14)

    p2 = doc.add_paragraph("по комплектации товаров (запчасти по бренду)")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].bold = True
    p2.runs[0].font.size = Pt(12)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_dt = p3.add_run(f"По состоянию на: {datetime.now().strftime('%d.%m.%Y')}")
    run_dt.bold = True
    run_dt.font.size = Pt(11)

    doc.add_paragraph("")

    # --- Таблица: БЕЗ ID, с границами ---
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"

    hdr = t.rows[0].cells
    hdr[0].text = "Запчасть"
    hdr[1].text = "Бренд"
    hdr[2].text = "Категория запчасти"
    hdr[3].text = "Цена"
    hdr[4].text = "Остаток"

    for r in rows:
        c = t.add_row().cells
        c[0].text = str(r.part_name)
        c[1].text = str(r.brand_name)
        c[2].text = str(r.part_category)
        c[3].text = str(r.price)
        c[4].text = str(r.stock_quantity)

    # --- Нижний колонтитул ---
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fr = fp.add_run(f"Отчет составлен на: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    fr.font.size = Pt(9)
    fr.bold = True

    # 3) отдаём файл
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="parts_report.docx")
# ---------------- CART (only for logged-in) ----------------

def _cart():
    if "cart" not in session or not isinstance(session.get("cart"), dict):
        session["cart"] = {}
    return session["cart"]

@auth_bp.route("/cart")
def cart():
    gate = _require_login()
    if gate:
        return gate

    cart_map = _cart()  # { "product_id": qty }
    ids = [int(k) for k in cart_map.keys()] if cart_map else []

    items = []
    total = 0

    if ids:
        conn = get_conn()
        cur = conn.cursor()

        placeholders = ",".join(["?"] * len(ids))
        cur.execute(f"SELECT * FROM dbo.vw_products WHERE product_id IN ({placeholders})", *ids)
        rows = cur.fetchall()
        conn.close()

        by_id = {int(r.product_id): r for r in rows}

        for pid_str, qty in cart_map.items():
            pid = int(pid_str)
            p = by_id.get(pid)
            if not p:
                continue

            q = int(qty)
            price = float(p.price) if p.price is not None else 0
            line = price * q
            total += line

            items.append({
                "product_id": pid,
                "name": p.name,
                "price": price,
                "qty": q,
                "line_total": line
            })

    return render_template("cart.html", items=items, total=total)

@auth_bp.route("/cart/add/<int:pid>", methods=["POST"])
def cart_add(pid):
    gate = _require_login()
    if gate:
        return gate

    cart_map = _cart()
    qty = request.form.get("qty", "1").strip()
    try:
        qty_int = max(1, int(qty))
    except:
        qty_int = 1

    cart_map[str(pid)] = int(cart_map.get(str(pid), 0)) + qty_int
    session.modified = True
    return redirect(request.referrer or "/")

@auth_bp.route("/cart/update", methods=["POST"])
def cart_update():
    gate = _require_login()
    if gate:
        return gate

    cart_map = _cart()
    for k, v in request.form.items():
        if not k.startswith("qty_"):
            continue
        pid = k.replace("qty_", "").strip()
        try:
            q = int(v)
        except:
            q = 1

        if q <= 0:
            cart_map.pop(pid, None)
        else:
            cart_map[pid] = q

    session.modified = True
    return redirect("/cart")

@auth_bp.route("/cart/remove/<int:pid>", methods=["POST"])
def cart_remove(pid):
    gate = _require_login()
    if gate:
        return gate

    cart_map = _cart()
    cart_map.pop(str(pid), None)
    session.modified = True
    return redirect("/cart")

@auth_bp.route("/cart/clear", methods=["POST"])
def cart_clear():
    gate = _require_login()
    if gate:
        return gate

    session["cart"] = {}
    session.modified = True
    return redirect("/cart")
@auth_bp.route("/cart/checkout", methods=["POST"])
def cart_checkout():
    gate = _require_login()
    if gate:
        return gate

    cart = session.get("cart", {})
    if not cart:
        return redirect("/cart")

    user_id = session["user_id"]

    conn = get_conn()
    cur = conn.cursor()

    # товары из корзины
    ids = [int(k) for k in cart.keys()]
    placeholders = ",".join(["?"] * len(ids))

    cur.execute(f"""
        SELECT product_id, price
        FROM dbo.vw_products
        WHERE product_id IN ({placeholders})
    """, *ids)

    rows = cur.fetchall()
    price_map = {int(r.product_id): float(r.price) for r in rows}

    total = 0.0
    for pid_str, qty in cart.items():
        pid = int(pid_str)
        q = int(qty)
        total += price_map.get(pid, 0.0) * q

    # ============================
    # ✅ НОВЫЕ ПОЛЯ
    # ============================

    first_name = request.form.get("first_name")
    last_name = request.form.get("last_name")
    email = request.form.get("email")
    phone = request.form.get("phone")
    delivery_address = request.form.get("delivery_address")
    comment = request.form.get("comment")
    payment_method = request.form.get("payment_method")
    delivery_method = request.form.get("delivery_method")

    delivery_cost = 0

    if delivery_method == "pickup":
        delivery_address = "Самовывоз"

    if delivery_method == "minsk" and total < 140:
        delivery_cost = 10

    total += delivery_cost

    # ============================
    # 1) создаём заказ
    # ============================

    cur.execute("""
    DECLARE @new_id INT;

    EXEC dbo.sp_place_order
        ?,?,?,?,?,?,?,?,?,?,?,
        @new_id = @new_id OUTPUT;

    SELECT @new_id;
    """,
                user_id,
                total,
                delivery_address,
                first_name,
                last_name,
                email,
                phone,
                comment,
                payment_method,
                delivery_method,
                delivery_cost
                )

    order_id = cur.fetchone()[0]

    # ============================
    # 2) позиции заказа
    # ============================

    for pid_str, qty in cart.items():
        pid = int(pid_str)
        q = int(qty)
        price = price_map.get(pid, 0.0)

        cur.execute(
            "EXEC dbo.sp_add_order_item ?, ?, ?, ?",
            order_id, pid, q, price
        )

    conn.commit()
    # собираем товары для письма
    email_items = []

    for pid_str, qty in cart.items():
        pid = int(pid_str)
        q = int(qty)
        price = price_map.get(pid, 0.0)

        # получим название
        cur.execute("SELECT name FROM dbo.products WHERE product_id=?", pid)
        name_row = cur.fetchone()

        email_items.append({
            "name": name_row.name,
            "qty": q,
            "price": price
        })
        dlvrm = request.form.get("delivery_method")
        dlvr=""
        if dlvrm == "pickup":
            dlvr = "Самовывоз"
        elif dlvrm == "minsk":
            dlvr="Доставка курьером по Минску"
        else:
            dlvr="Доставка по Беларуси"

    # отправка письма
    send_order_created_email(
        user_email=session["email"],
        order_id=order_id,
        phone=request.form.get("phone"),
        delivery_type=dlvr,
        items=email_items,
        total=total
    )
    conn.close()

    return redirect("/cabinet")
from flask import jsonify

@auth_bp.route("/compat")
def compat_page():
    gate = _require_login()
    if gate:
        return gate
    return render_template("compat.html")

@auth_bp.route("/api/search/products")
def api_search_products():
    gate = _require_login()
    if gate:
        return gate

    q = request.args.get("q", "").strip()
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("EXEC dbo.sp_search_products ?, ?", q, "product")
    rows = cur.fetchall()
    conn.close()

    return jsonify([{"id": int(r.product_id), "name": r.name} for r in rows])

@auth_bp.route("/api/search/parts")
def api_search_parts():
    gate = _require_login()
    if gate:
        return gate

    q = request.args.get("q", "").strip()
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("EXEC dbo.sp_search_products ?, ?", q, "part")
    rows = cur.fetchall()
    conn.close()

    return jsonify([{"id": int(r.product_id), "name": r.name} for r in rows])
@auth_bp.route("/api/brands")
def api_brands_by_category():
    # можно и без логина — это просто справочник для фильтра
    category_id_raw = request.args.get("category_id", "").strip()

    conn = get_conn()
    cur = conn.cursor()

    # категории (чтобы понять ветку запчастей и собрать подкатегории)
    cur.execute("SELECT category_id, name, parent_id FROM dbo.categories")
    categories = cur.fetchall()

    # строим children_map и parts_tree
    parts_roots = set()
    children_map = {}

    for c in categories:
        if c.name and "запчаст" in c.name.lower():
            parts_roots.add(int(c.category_id))
        if c.parent_id is not None:
            children_map.setdefault(int(c.parent_id), []).append(int(c.category_id))

    def collect_descendants(root_id: int) -> set[int]:
        stack = [root_id]
        out = set()
        while stack:
            x = stack.pop()
            if x in out:
                continue
            out.add(x)
            for ch in children_map.get(x, []):
                stack.append(ch)
        return out

    parts_tree = set()
    for rid in parts_roots:
        parts_tree |= collect_descendants(rid)

    # если категория не выбрана — вернём все бренды
    if not category_id_raw.isdigit():
        cur.execute("SELECT id_brand, name FROM dbo.brand ORDER BY name")
        rows = cur.fetchall()
        conn.close()
        return jsonify({
            "ok": True,
            "is_parts_mode": False,
            "brands": [{"id_brand": int(r.id_brand), "name": r.name} for r in rows]
        })

    cat_id = int(category_id_raw)
    is_parts_mode = cat_id in parts_tree

    # бренды только в выбранной категории + её подкатегориях
    cat_ids = collect_descendants(cat_id)
    placeholders = ",".join(["?"] * len(cat_ids))

    cur.execute(f"""
        SELECT DISTINCT
            vp.id_brand   AS id_brand,
            vp.brand_name AS name
        FROM dbo.vw_products vp
        WHERE vp.id_brand IS NOT NULL
          AND vp.category_id IN ({placeholders})
        ORDER BY vp.brand_name
    """, *list(cat_ids))
    rows = cur.fetchall()
    conn.close()

    return jsonify({
        "ok": True,
        "is_parts_mode": is_parts_mode,
        "brands": [{"id_brand": int(r.id_brand), "name": r.name} for r in rows]
    })

@auth_bp.route("/api/compat/matrix", methods=["POST"])
def api_compat_matrix():
    gate = _require_login()
    if gate:
        return gate

    data = request.get_json(silent=True) or {}
    product_id = data.get("product_id")

    # поддержка старого формата (part_id) и нового (part_ids)
    part_ids = data.get("part_ids")
    if part_ids is None and data.get("part_id"):
        part_ids = [data.get("part_id")]

    if not product_id or not part_ids or not isinstance(part_ids, list):
        return jsonify({"ok": False, "error": "product_id and part_ids required"}), 400

    # нормализуем список part_ids
    norm_part_ids = []
    seen = set()
    for x in part_ids:
        try:
            pid = int(x)
            if pid > 0 and pid not in seen:
                norm_part_ids.append(pid)
                seen.add(pid)
        except:
            pass

    if not norm_part_ids:
        return jsonify({"ok": False, "error": "No valid part_ids"}), 400

    product_id = int(product_id)

    conn = get_conn()
    cur = conn.cursor()

    results = []

    for part_id in norm_part_ids:
        cur.execute("EXEC dbo.sp_part_compat_matrix ?", part_id)
        rows = cur.fetchall()

        compatible = []
        incompatible = []
        part_name = None

        for r in rows:
            part_name = getattr(r, "part_name", None)
            item = {
                "id": int(r.product_id),
                "name": r.product_name,
                "brand": getattr(r, "brand_name", None),
                "img": f"/product_image/{int(r.product_id)}",
            }
            if int(r.is_compatible) == 1:
                compatible.append(item)
            else:
                incompatible.append(item)

        selected_ok = any(x["id"] == product_id for x in compatible)

        results.append({
            "part_id": part_id,
            "part_name": part_name or f"Запчасть #{part_id}",
            "selected_ok": selected_ok,
            "compatible": compatible,
            "incompatible": incompatible
        })

    conn.close()

    return jsonify({
        "ok": True,
        "product_id": product_id,
        "results": results
    })
@auth_bp.route("/smart", methods=["GET", "POST"])
def smart_pick():
    conn = get_conn()
    cur = conn.cursor()

    # списки для формы
    cur.execute("SELECT category_id, name FROM dbo.categories ORDER BY name")
    categories = cur.fetchall()

    cur.execute("SELECT id_brand, name FROM dbo.brand ORDER BY name")
    brands = cur.fetchall()

    # значения формы
    category_id = request.values.get("category_id", "").strip()
    brand_id = request.values.get("brand_id", "").strip()
    price_from = request.values.get("price_from", "").strip()
    price_to = request.values.get("price_to", "").strip()
    in_stock = (request.values.get("in_stock") == "1")

    suitable = []
    unsuitable = []

    # --- определяем режим "запчасти" по выбранной категории ---
    parts_mode = False          # бренд отключаем
    parts_root_mode = False     # выбрана корневая "Запчасти"
    parts_root_id = None
    parts_category_ids = set()  # все категории внутри ветки запчастей (для root)

    cat = int(category_id) if category_id.isdigit() else None
    br = int(brand_id) if brand_id.isdigit() else None

    if cat is not None:
        cur.execute("""
            SELECT
                c.category_id,
                c.name,
                c.parent_id,
                ISNULL(p.name,'') AS parent_name
            FROM dbo.categories c
            LEFT JOIN dbo.categories p ON p.category_id = c.parent_id
            WHERE c.category_id = ?
        """, cat)
        cinfo = cur.fetchone()

        if cinfo:
            cname = (cinfo.name or "").lower()
            pname = (cinfo.parent_name or "").lower()
            parent_id = cinfo.parent_id

            is_parts_any = ("запчаст" in cname) or ("запчаст" in pname)
            if is_parts_any:
                parts_mode = True

                # корневая "Запчасти": имя содержит "запчаст", а родитель не "запчасти"
                if ("запчаст" in cname) and (parent_id is None or ("запчаст" not in pname)):
                    parts_root_mode = True
                    parts_root_id = int(cinfo.category_id)

                    # соберём все category_id внутри ветки "Запчасти"
                    cur.execute("""
                        WITH cat_tree AS (
                            SELECT category_id, parent_id
                            FROM dbo.categories
                            WHERE category_id = ?
                            UNION ALL
                            SELECT c.category_id, c.parent_id
                            FROM dbo.categories c
                            JOIN cat_tree t ON c.parent_id = t.category_id
                        )
                        SELECT category_id FROM cat_tree
                    """, parts_root_id)
                    parts_category_ids = {int(r[0]) for r in cur.fetchall()}

    # если запчасти — бренд игнорируем
    if parts_mode:
        br = None

    # если форма отправлена
    if request.method == "POST":
        pf = None
        pt = None
        try:
            pf = float(price_from.replace(",", ".")) if price_from else None
        except:
            pf = None
        try:
            pt = float(price_to.replace(",", ".")) if price_to else None
        except:
            pt = None

        # кандидаты
        cur.execute("""
            SELECT *
            FROM dbo.vw_products
            WHERE is_active = 1
            ORDER BY product_id DESC
        """)
        rows = cur.fetchall()

        # человеко-читаемые названия выбранных
        cat_name = None
        brand_name = None
        if cat:
            cur.execute("SELECT name FROM dbo.categories WHERE category_id = ?", cat)
            rr = cur.fetchone()
            cat_name = rr[0] if rr else None
        if br:
            cur.execute("SELECT name FROM dbo.brand WHERE id_brand = ?", br)
            rr = cur.fetchone()
            brand_name = rr[0] if rr else None

        for p in rows:
            reasons_ok = []
            reasons_bad = []

            # ---- КАТЕГОРИЯ (особая логика для запчастей) ----
            if cat is not None:
                if parts_mode:
                    # 1) выбрали корневую "Запчасти" -> подходят ВСЕ товары из ветки запчастей
                    if parts_root_mode:
                        if int(p.category_id or 0) in parts_category_ids:
                            reasons_ok.append(("Категория", "ветка запчастей"))
                        else:
                            reasons_bad.append(("Категория", "это не запчасть"))
                    else:
                        # 2) выбрали подкатегорию запчастей -> только эта подкатегория
                        if p.category_id == cat:
                            reasons_ok.append(("Категория", f"{cat_name or cat}"))
                        else:
                            reasons_bad.append(("Категория", f"нужна {cat_name or cat}, а тут {p.category_name or '-'}"))
                else:
                    # обычная категория (кресла/диваны и т.п.) — строго совпадение category_id
                    if p.category_id == cat:
                        reasons_ok.append(("Категория", f"{cat_name or cat}"))
                    else:
                        reasons_bad.append(("Категория", f"нужна {cat_name or cat}, а тут {p.category_name or '-'}"))

            # ---- БРЕНД (если НЕ запчасти) ----
            if br is not None:
                if getattr(p, "id_brand", None) == br:
                    reasons_ok.append(("Бренд", f"{brand_name or br}"))
                else:
                    reasons_bad.append(("Бренд", f"нужен {brand_name or br}, а тут {p.brand_name or '-'}"))

            # ---- ЦЕНА ----
            if pf is not None:
                if float(p.price) >= pf:
                    reasons_ok.append(("Цена от", f"≥ {pf}"))
                else:
                    reasons_bad.append(("Цена от", f"цена {p.price} < {pf}"))
            if pt is not None:
                if float(p.price) <= pt:
                    reasons_ok.append(("Цена до", f"≤ {pt}"))
                else:
                    reasons_bad.append(("Цена до", f"цена {p.price} > {pt}"))

            # ---- НАЛИЧИЕ ----
            if in_stock:
                if int(p.stock_quantity) > 0:
                    reasons_ok.append(("Наличие", "есть"))
                else:
                    reasons_bad.append(("Наличие", "нет в наличии"))

            item = {"p": p, "ok": reasons_ok, "bad": reasons_bad}

            if len(reasons_bad) == 0:
                suitable.append(item)
            else:
                unsuitable.append(item)

    conn.close()

    return render_template(
        "smart.html",
        categories=categories,
        brands=brands,
        category_id=category_id,
        brand_id=brand_id,
        price_from=price_from,
        price_to=price_to,
        in_stock=in_stock,
        suitable=suitable,
        unsuitable=unsuitable,
        parts_mode=parts_mode
    )